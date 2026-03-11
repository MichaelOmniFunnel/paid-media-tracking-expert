"""
OFM Repo Audit: Consolidated Word Document Builder
Generates a professional Word document covering both Batch 1 (49 repos) and Batch 2 (33 items)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import datetime

doc = Document()

# -- Page Setup --
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# -- Styles --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

# Helper functions
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    return h

def add_verdict_block(number, name, verdict, priority, why, application, conflicts):
    p = doc.add_paragraph()
    run = p.add_run(f"{number}. {name}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    # Verdict line with color coding
    vp = doc.add_paragraph()
    vr = vp.add_run("Verdict: ")
    vr.bold = True
    vr.font.size = Pt(10)
    vv = vp.add_run(verdict)
    vv.bold = True
    vv.font.size = Pt(10)
    if "ADOPT" in verdict.upper():
        vv.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    elif "SKIP" in verdict.upper() or "HARD SKIP" in verdict.upper():
        vv.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    elif "STUDY" in verdict.upper():
        vv.font.color.rgb = RGBColor(0xCC, 0x88, 0x00)
    elif "MONITOR" in verdict.upper():
        vv.font.color.rgb = RGBColor(0x66, 0x66, 0xCC)
    elif "ALREADY" in verdict.upper():
        vv.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    vp.add_run(f"  |  Priority: ").bold = True
    vp.add_run(priority)

    # Why
    wp = doc.add_paragraph()
    wr = wp.add_run("Why: ")
    wr.bold = True
    wp.add_run(why)

    # Application
    ap = doc.add_paragraph()
    ar = ap.add_run("Application: ")
    ar.bold = True
    ap.add_run(application)

    # Conflicts
    cp = doc.add_paragraph()
    cr = cp.add_run("Conflicts: ")
    cr.bold = True
    cp.add_run(conflicts)

    # Spacer
    doc.add_paragraph()

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()
    return table


# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_heading('OmniFunnel Marketing', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    run.font.size = Pt(28)

subtitle = doc.add_heading('Claude Code Repository Audit', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.color.rgb = RGBColor(0x33, 0x66, 0x99)
    run.font.size = Pt(20)

sub2 = doc.add_heading('Consolidated Assessment: Batch 1 + Batch 2', level=2)
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in sub2.runs:
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f"Date: {datetime.date.today().strftime('%B %d, %Y')}\n").bold = True
meta.add_run("Prepared by: Claude Code (Senior Technical Architect)\n")
meta.add_run("For: Michael Tate, OmniFunnel Marketing\n")
meta.add_run("Status: DRAFT\n").bold = True
meta.add_run(f"\nBatch 1: 49 repositories across 6 categories\n")
meta.add_run(f"Batch 2: 33 items across 10 categories\n")
meta.add_run(f"Total: 82 items evaluated")

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
add_heading_styled("Table of Contents", 1)
toc_items = [
    "1. Executive Summary",
    "2. Scorecard",
    "3. Batch 1: Individual Reviews (49 repos)",
    "   3.1 Context Management & Session Persistence",
    "   3.2 Behavioral Frameworks & Cognitive Architecture",
    "   3.3 Claude Code Infrastructure & Toolkits",
    "   3.4 Subagent Collections & Orchestration",
    "   3.5 Paid Media & Marketing Specific",
    "   3.6 Development & Browser Automation",
    "4. Batch 1: Top 10 Power Stack",
    "5. Batch 1: Architecture Blueprint",
    "6. Batch 1: Custom Build Recommendations",
    "7. Batch 2: Individual Reviews (33 items)",
    "   7.1 Project Execution & Workflow Engines",
    "   7.2 Context Engineering & Architecture Knowledge",
    "   7.3 Skill Generation & Management",
    "   7.4 Web Research & Scraping",
    "   7.5 Prompt Preprocessing & Input Quality",
    "   7.6 Content Quality & Writing",
    "   7.7 Security & Configuration",
    "   7.8 Monitoring, Directories & Ecosystem Tracking",
    "   7.9 Social Media & Publishing",
    "   7.10 Upcoming / Just Announced",
    "8. Batch 2: Top 5",
    "9. Combined Power Stack (15 repos, both batches)",
    "10. Redundancy Report",
    "11. Custom Build Recommendations (Batch 2)",
    "12. Implementation Action Plan",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

doc.add_page_break()

# ============================================================
# EXECUTIVE SUMMARY
# ============================================================
add_heading_styled("1. Executive Summary", 1)

doc.add_paragraph(
    "This document is the consolidated assessment of 82 GitHub repositories, tools, and upcoming Anthropic features "
    "evaluated for OmniFunnel Marketing's Claude Code infrastructure. The audit was conducted across two batches: "
    "Batch 1 (49 repos covering MCP servers, paid media tools, subagent collections, and foundational infrastructure) "
    "and Batch 2 (33 items covering context engineering, skill generation, web research, security, content quality, "
    "and upcoming platform features)."
)

doc.add_paragraph(
    "OFM's existing setup is already sophisticated: 27 custom skills, 9 domain specific agents, 6 hooks, "
    "6 MCP servers (Chrome, Meta Ads, Asana, GA4, GSC, Google Ads), a client memory protocol, and comprehensive "
    "rules for security, anomaly flagging, and verification. The majority of repos evaluated (42 of 82) were rated "
    "SKIP because they either solve problems OFM does not have, would conflict with existing infrastructure, or "
    "add complexity without proportional value."
)

doc.add_paragraph(
    "The repos that matter fall into three categories: (1) tools that fill genuine gaps in OFM's current capabilities "
    "(Context7, Firecrawl, Growth Marketing Agents, AgentShield), (2) reference materials that make existing "
    "infrastructure smarter (Claude Code System Prompts, Meta Ads Analyzer references, Hooks Mastery), and "
    "(3) upcoming Anthropic features that will transform OFM's operational model (Auto Mode, /loop command). "
    "The biggest competitive advantage, however, comes from 7 custom builds that no open source repo addresses: "
    "client context flags, Stape SGTM monitoring, MCP response sandboxing, weighted audit scoring, cross platform "
    "anomaly dashboards, Klaviyo/ActiveCampaign integration, and NetSuite/SuiteCommerce tracking."
)

doc.add_page_break()

# ============================================================
# SCORECARD
# ============================================================
add_heading_styled("2. Scorecard", 1)

add_heading_styled("Batch 1 (49 repos)", 2)
add_table(
    ["Category", "Count"],
    [
        ["ADOPT", "5"],
        ["ALREADY ADOPTED", "2"],
        ["STUDY", "21"],
        ["SKIP", "21"],
    ]
)

add_heading_styled("Batch 2 (33 items)", 2)
add_table(
    ["Category", "Count"],
    [
        ["ADOPT", "5 (includes Auto Mode and /loop)"],
        ["STUDY", "10"],
        ["MONITOR", "3"],
        ["SKIP", "14"],
        ["HARD SKIP", "1"],
    ]
)

add_heading_styled("Combined (82 items)", 2)
add_table(
    ["Category", "Count"],
    [
        ["ADOPT (install/enable)", "10"],
        ["ALREADY ADOPTED", "2"],
        ["STUDY (extract patterns)", "31"],
        ["MONITOR (watch for maturity)", "3"],
        ["SKIP", "35"],
        ["HARD SKIP", "1"],
    ]
)

doc.add_page_break()

# ============================================================
# BATCH 1: INDIVIDUAL REVIEWS
# ============================================================
add_heading_styled("3. Batch 1: Individual Reviews (49 repos)", 1)

# -- Category 1 --
add_heading_styled("3.1 Context Management & Session Persistence", 2)

add_verdict_block(1, "Continuous Claude v3",
    "SKIP", "N/A",
    "Requires PostgreSQL with pgvector, Docker, Python 3.11+, FAISS vector embeddings, and a daemon process. Installs 109 skills, 32 agents, and 30 hooks that would obliterate OFM's 27 paid media skills, 9 domain agents, and 6 hooks. The 'compound, don't compact' philosophy solves a problem that matters more for long running dev projects than discrete paid media client engagements.",
    "N/A. OFM's existing client memory protocol, MEMORY.md, and PreCompact hook already provide session continuity tailored to paid media.",
    "Catastrophic. Would overwrite every custom skill, agent, and hook OFM has built.")

add_verdict_block(2, "Context Mode",
    "STUDY", "P2",
    "Two genuinely valuable innovations. First, the MCP virtualization layer runs commands in sandboxed environments where only stdout enters context, achieving 99%+ token reduction on large outputs. This addresses OFM's problem with verbose GAQL results and Meta insights breakdowns consuming context. Second, the PreCompact hook generates a structured 2KB XML snapshot with 15 categories, far more sophisticated than OFM's current plaintext preservation.",
    "(1) Build a sandbox wrapper for OFM's most verbose MCP calls (Google Ads GAQL, Meta insights with daily breakdowns). (2) Upgrade pre-compact-preserve.js to capture structured snapshot with active client, audit phase, findings, decisions, and files modified.",
    "PreCompact and SessionStart hooks collide with OFM's existing hooks. Cherry pick the concepts, do not install the hooks.")

add_verdict_block(3, "Claude Mem",
    "STUDY", "P3",
    "The standout concept is progressive disclosure for memory retrieval: inject 50 to 100 token summaries first, expand to full entries only when needed. OFM currently reads entire client history.md files on re-engagement, which becomes wasteful as files grow past 200+ lines. However, the implementation requires Bun runtime, Chroma vector database, and 6 lifecycle hooks that conflict with OFM's infrastructure. AGPL-3.0 license restricts derivative use.",
    "Restructure client history.md files with summary headers (last engagement, current status, active campaigns, open items count) followed by collapsible sections. Pattern to adopt, not tool to install.",
    "High. Six hooks collide with OFM's existing hooks. AGPL license problematic for proprietary use.")

add_verdict_block(4, "Claude Context Monitor",
    "SKIP", "N/A",
    "macOS only. Uses osascript (AppleScript), macOS specific stat flags, and macOS notification center. OFM runs on Windows 11. Zero stars confirms this is a personal utility with no cross platform support.",
    "N/A. Platform incompatible.",
    "Cannot run on Windows.")

add_verdict_block(5, "Claude Session Restore",
    "SKIP", "N/A",
    "7 stars, requires Rust 1.70+ to compile. OFM already handles session continuity through 4 independent mechanisms (client memory files, MEMORY.md, PreCompact hook, Asana task state). Parsing raw JSONL session files after the fact is less efficient than capturing structured context at write time.",
    "N/A.",
    "Adds Rust build toolchain dependency for marginal benefit.")

# -- Category 2 --
add_heading_styled("3.2 Behavioral Frameworks & Cognitive Architecture", 2)

add_verdict_block(6, "SuperClaude Framework",
    "SKIP", "N/A",
    "16 generic agents (PM Agent, Security Engineer, Frontend Architect) replace domain expertise with generic dev capabilities. OFM's 9 agents know ROAS, CAPI, GTM, bid strategies. SuperClaude's know nothing about paid media. The behavioral trait system (ULTRA_DETAIL, SPEED_DEMON) conflicts with OFM's precisely calibrated CLAUDE.md communication rules.",
    "N/A.",
    "Severe. 16 agents conflict with 9 domain agents. 30 commands conflict with 27 skills. Behavioral traits override CLAUDE.md rules.")

add_verdict_block(7, "Superpowers",
    "SKIP", "N/A",
    "76K stars, but enforces mandatory TDD with RED/GREEN/REFACTOR cycles, forced brainstorming gates, and automated code review. Excellent for building software. Irrelevant and actively harmful for paid media audits, tracking implementation, and campaign analysis. 'Zero config' means these behaviors activate automatically and cannot be selectively disabled.",
    "N/A. No paid media workflow benefits from mandatory TDD.",
    "Pervasive. Auto triggered skills impose software engineering ceremony on every task.")

add_verdict_block(8, "GSD (Get Shit Done)",
    "STUDY", "P2",
    "Treats the 200K context window as a finite, managed resource. Phase based workflow (discuss, plan, execute, verify) where each execution wave gets fresh context. The 47th task executes with the same quality as the 1st because it is not carrying accumulated context garbage. Multi agent orchestration with fresh context per executor addresses context bleed during multi client audit days.",
    "(1) Adapt phase model as a new OFM skill for complex engagements (full tracking implementations, cross platform audits). (2) Fresh context per execution wave eliminates cross client contamination during multi client audit days. (3) The verify phase maps to OFM's Asana verification comment workflow.",
    "Low. .planning/ directory is self contained, does not touch .claude/. Study patterns and build OFM specific skill.")

add_verdict_block(9, "Everything Claude Code",
    "STUDY", "P1",
    "Most mature and relevant behavioral repo. Three novel features: (1) Instinct based learning with confidence scoring formalizes OFM's ad hoc self learning protocol into a quantified pipeline. (2) Hook profiles (env variable controls minimal/standard/strict) let OFM run lighter on simple reads and heavier on complex audits. (3) Checkpoint to external systems for phase transition documentation aligns with Asana verification workflow. Anthropic hackathon winner, 10+ months daily production use.",
    "Build three OFM native features: instinct pipeline in .claude/memory/instincts.md, hook profiles in settings.local.json, and checkpoint to Asana integration. Extract concepts, do not install ECC's agents/hooks/skills.",
    "Medium but modular. ECC's 9 generic agents, 10+ hooks, and 11 dev skills conflict with OFM's infrastructure. But the ideas are extractable without importing any code.")

# -- Category 3 --
add_heading_styled("3.3 Claude Code Infrastructure & Toolkits", 2)

add_verdict_block(10, "Awesome Claude Code",
    "SKIP", "N/A",
    "Curated list of links. OFM already operates deep inside the Claude Code ecosystem. A directory of links adds zero operational value when you already know the landscape and have purpose built tooling.",
    "N/A.",
    "None, but no substance.")

add_verdict_block(11, "Claude Code Showcase",
    "STUDY", "P2",
    "The skill evaluation system (UserPromptSubmit hook that auto suggests skills based on prompt keywords) is genuinely useful. OFM has 27 skills and Michael must remember which to invoke. Auto suggestion eliminates that cognitive load. GitHub Actions for automated docs sync and code quality sweeps are solid CI/CD patterns.",
    "Adapt the skill eval hook so Claude auto suggests relevant OFM skills based on prompt keywords. The JIRA ticket command pattern could inform an Asana equivalent.",
    "Would be a new UserPromptSubmit hook that complements existing hooks.")

add_verdict_block(12, "Awesome Claude Code Toolkit",
    "STUDY", "P3",
    "135 agents and 121 plugins, all generic dev focused. None target paid media. However, the plugin marketplace architecture and selective installation pattern is worth studying as OFM scales past 30+ skills.",
    "Study organizational structure for future scaling.",
    "Installing wholesale would flood OFM with 135 irrelevant agents.")

add_verdict_block(13, "Claude Code Hooks Mastery",
    "ADOPT", "P1",
    "Most comprehensive hooks reference with all 13 hook events implemented in working code. OFM currently uses only 2 hook events (PreToolUse, PreCompact). This demonstrates 11 additional hooks OFM is not leveraging: SessionStart (auto load context), Stop (session summaries), SubagentStop (monitor agents), Notification (alerts), PermissionRequest (auto allow read only). The UV single file script architecture is cleaner than OFM's current JS approach.",
    "(1) SessionStart hook to auto load client context and check Asana board. (2) Stop hook for auto session summaries. (3) PermissionRequest to auto approve read only MCP ops. (4) SubagentStart/SubagentStop for swarm monitoring. (5) Meta agent pattern for generating client specific agents.",
    "OFM's hooks are JS, this uses Python with UV. Can coexist since hooks support mixed languages.")

add_verdict_block(14, "Agents (wshobson)",
    "STUDY", "P2",
    "Three tier model strategy (Opus/Sonnet/Haiku for cost optimization) and 16 workflow orchestrators represent mature architecture. Includes 4 marketing plugins (SEO content, technical SEO) and 3 business plugins (analytics, operations). The progressive disclosure for skills and orchestrator patterns could improve OFM's multi agent delegation.",
    "Study workflow orchestrator patterns, three tier model routing for cost optimization, and plugin architecture as organizational model.",
    "Agents need reworking for OFM's safety model (no platform writes).")

add_verdict_block(15, "Agentwise",
    "SKIP", "N/A",
    "Full application development platform (335K+ lines). Solves a completely different problem. The token optimization only applies to its own multi agent system, not Claude Code's native agents.",
    "N/A.",
    "Entirely different orchestration model.")

add_verdict_block(16, "Claude Code Skill Factory",
    "STUDY", "P3",
    "Factory approach for generating skills, agents, hooks at scale. 69 prompt presets. If OFM scales to 50+ skills or starts templating client specific configurations, the factory pattern could reduce creation time.",
    "Low urgency. Review hooks factory with safety validation against current process.",
    "Uses own YAML conventions that differ from OFM's SKILL.md format.")

add_verdict_block(17, "Claude Code Guide",
    "SKIP", "N/A",
    "Documentation, not tooling. OFM has working implementations of everything this guide covers.",
    "N/A.",
    "None.")

add_verdict_block(18, "Claude Code Best Practice",
    "STUDY", "P2",
    "13K stars. Three immediately relevant patterns: (1) '/compact at 50%' rule for context window management. (2) 'Feature specific subagents over generic' validates OFM's agent approach. (3) 'Command to Agent to Skill' orchestration architecture formalizes OFM's workflow. Tips from Claude's creators may contain non obvious optimizations.",
    "Implement 50% compact discipline. Review agents against 'feature specific' pattern. Study orchestration architecture.",
    "The CLAUDE.md under 200 lines recommendation may conflict with OFM's current length. Worth evaluating.")

add_verdict_block(19, "Anthropic Official Skills",
    "ADOPT", "P1",
    "Anthropic's own skills repository. Document creation skills (DOCX, PPTX, PDF, XLSX) directly solve a recurring OFM need: all deliverables are Word documents, audits need professional formatting, QBRs need presentations. Production grade, Anthropic maintained. The spec/ directory defines the canonical skill standard OFM's skills should follow.",
    "(1) Install document skills immediately for DOCX deliverables. (2) XLSX for budget reports and performance exports. (3) PPTX for quarterly reviews. (4) PDF for form extraction. (5) Validate OFM's 27 skills against canonical spec.",
    "OFM's SKILL.md format appears to align with Anthropic's structure. Installation should be clean.")

add_verdict_block(20, "Awesome Agent Skills (VoltAgent)",
    "STUDY", "P3",
    "549+ skills from real engineering teams (Anthropic, Google Labs, Vercel, Stripe, Cloudflare). None target paid media. Value is as a reference catalog: when OFM needs a specific platform integration, check here first for an official skill.",
    "Bookmark as reference. Not actionable today.",
    "None. Catalog, not installation.")

# -- Category 4 --
add_heading_styled("3.4 Subagent Collections & Orchestration", 2)

add_verdict_block(21, "Awesome Claude Code Subagents (VoltAgent)",
    "SKIP", "N/A",
    "127+ subagents, virtually all software development focused. Zero paid media, marketing, or advertising agents. OFM already has 9 purpose built agents.",
    "N/A.",
    "Generic orchestration would create confusion alongside OFM's custom system.")

add_verdict_block(22, "Claude Code Subagents (0xfurai)",
    "SKIP", "N/A",
    "100+ agents, all software engineering focused (React, Docker, Kubernetes). Zero marketing or analytics agents. Developer toolkit, not a paid media toolkit.",
    "N/A.",
    "Would pollute OFM's curated agent directory.")

add_verdict_block(23, "Claude Code Sub Agents (lst97)",
    "SKIP", "N/A",
    "33 agents, entirely software development. The agent organizer meta orchestrator is already handled by OFM's CLAUDE.md. The cost awareness note about 2 to 5x token usage for multi agent workflows is worth remembering but does not justify adoption.",
    "N/A.",
    "Agent organizer competes with OFM's existing orchestration logic.")

add_verdict_block(24, "Awesome Claude Agents (rahulvrane)",
    "STUDY", "P3",
    "Curated ecosystem directory, not installable code. Value as a periodic discovery tool to scan for new marketing relevant agents or patterns. Not something to install.",
    "Use as quarterly discovery feed. Scan for marketing specific entries.",
    "None. Reference list.")

add_verdict_block(25, "Claude Subagent System",
    "STUDY", "P3",
    "Two patterns worth studying: (1) hard token limits per agent (400 to 1500 tokens) to prevent scope creep and control costs, (2) quality gates at stage transitions that verify output completeness before next stage begins. Domain (product development) is completely wrong.",
    "Quality gate pattern could verify audit completeness before synthesis. Token limits could control swarm mode costs.",
    "Hierarchical authority model adds unnecessary complexity on top of what works.")

# -- Category 5 --
add_heading_styled("3.5 Paid Media & Marketing Specific", 2)

add_verdict_block(26, "Claude Ads",
    "STUDY", "P1",
    "Most relevant single repo. 190 audit checks across Google (74), Meta (46), LinkedIn (25), TikTok (25), Microsoft (20). Severity weighted scoring produces A through F grades. 11 industry templates covering ecommerce, legal, healthcare, finance. OFM already has its own audit methodology, but the structured check inventory is the real value as a gap analysis source.",
    "(1) Cross reference 190 checks against OFM's existing audit skills to find gaps. (2) Evaluate weighted A through F scoring for client deliverables. (3) Review industry templates against OFM's verticals. (4) LinkedIn + Microsoft checks provide foundation if those platforms enter scope.",
    "Heavy overlap with OFM's full audit, tracking audit, creative audit skills and all 9 agents. Treat as reference mine, not replacement.")

add_verdict_block(27, "Ads MCP (Adspirer)",
    "STUDY", "P2",
    "Cross platform MCP with LinkedIn Ads (28 tools) and TikTok Ads (4 tools), platforms where OFM currently relies on Chrome only. The strategy persistence pattern (saving decisions to STRATEGY.md) is novel. However, Google and Meta tools overlap with OFM's existing MCP servers.",
    "If OFM needs programmatic LinkedIn Ads access, this is the source. Adapt strategy persistence pattern for client memory as clients/[name]/strategy.md.",
    "Overlaps with Google Ads MCP and Meta Ads MCP. Never install alongside existing servers.")

add_verdict_block(28, "Google Ads MCP (cohnen)",
    "ALREADY ADOPTED", "N/A",
    "This is the exact MCP server OFM already runs. 5 tools, read only, OAuth, GAQL support. Already configured with MCC 677-900-1476 across 19 accounts.",
    "Already the backbone of OFM's Google Ads data layer.",
    "N/A. Incumbent.")

add_verdict_block(29, "Google Ads MCP Server (bjorndavidhansen)",
    "STUDY", "P2",
    "Goes beyond cohnen's with analytical layers: budget analysis (utilization rates), search term analysis, statistical anomaly detection (period over period with deviation thresholds), and automated insights (bid recommendations, negative keyword suggestions). The anomaly detection aligns directly with OFM's anomaly-flagging.md.",
    "Study anomaly detection algorithms for encoding into OFM's keyword strategist and budget optimizer agents. Study budget analysis module. Do NOT replace cohnen's server.",
    "Direct competitor to cohnen. Write capabilities violate OFM's read only safety model. Study logic only.")

add_verdict_block(30, "Google Ads MCP (gomarble)",
    "SKIP", "N/A",
    "Same core functionality as cohnen plus keyword planner. Not worth migration risk for one incremental tool. OFM's keyword strategist already handles keyword research through GAQL and Chrome.",
    "N/A.",
    "Direct duplicate of OFM's existing Google Ads MCP.")

add_verdict_block(31, "TrueClicks Google Ads MCP",
    "SKIP", "N/A",
    "Archived January 2026. Google released official MCP. Dead project. However, the signal that Google has an official MCP is strategically important. OFM should evaluate the official Google Ads MCP as a long term cohnen replacement.",
    "N/A. Monitor Google's official MCP for maturity.",
    "Archived.")

add_verdict_block(32, "Meta Ads MCP (Pipeboard)",
    "ALREADY ADOPTED", "N/A",
    "This is the exact MCP server OFM already runs. 24+ tools, remote hosted, 27 ad accounts connected.",
    "Already the backbone of OFM's Meta Ads data layer.",
    "N/A. Incumbent.")

add_verdict_block(33, "Meta Ads Analyzer",
    "ADOPT", "P1",
    "Fills a genuine knowledge gap with 9 reference documents on Meta's internal mechanics: Breakdown Effect (why the algorithm allocates budget to 'bad looking' segments), Learning Phase diagnostics with resolution steps, Auction Overlap analysis, Pacing Mechanics, and Creative Fatigue detection beyond simple frequency metrics. These are the exact explanations OFM has to improvise for clients repeatedly.",
    "(1) Extract 9 reference docs into .claude/skills/meta-ads-tracking/references/. (2) Add Breakdown Effect to client reporting templates. (3) Add Learning Phase checklist to platform strategist and creative analyst agents. (4) Creative Fatigue framework complements creative analyst agent. Do NOT install the MCP server (Pipeboard is already in production).",
    "MCP server overlaps with Pipeboard. Only adopt the skill layer and reference documents.")

add_verdict_block(34, "Meta MCP (brijr)",
    "SKIP", "N/A",
    "OFM already has Pipeboard. This has only 101 stars, last updated January 2025 (stale). No additive value.",
    "N/A.",
    "Direct overlap with Pipeboard.")

add_verdict_block(35, "Google Analytics 4 MCP",
    "SKIP", "N/A",
    "OFM already runs GA4 MCP covering 34 properties via OAuth. This uses service account auth (less flexible for multi client) and has only 186 stars with January 2025 last update.",
    "N/A.",
    "Direct overlap with existing GA4 MCP.")

add_verdict_block(36, "Google Search Console MCP",
    "STUDY", "P3",
    "521 stars, actively maintained, 19 tools. Notable features OFM's current GSC MCP may lack: compare_search_periods, batch_url_inspection, get_advanced_search_analytics. Worth comparing tool for tool.",
    "If current GSC MCP lacks batch inspection or period comparison, consider swapping.",
    "Overlaps with existing GSC MCP.")

add_verdict_block(37, "Marketing Skills",
    "STUDY", "P2",
    "12.6K stars, 34 skills across CRO, copywriting, SEO, analytics, paid advertising, email, landing pages. The value is cherry picking frameworks for areas where OFM's skills are thinner (content/copy, churn prevention, referral programs, pricing strategy). Hub and spoke model with product marketing context as foundation is an interesting architecture.",
    "Review paid advertising, analytics, and landing page skills against OFM's equivalents. Extract superior frameworks or prompt engineering patterns.",
    "Overlaps with conversion-optimization, landing-page-audit, creative-audit, ab-testing skills. Generalist approach versus OFM's agency specific approach. Complementary.")

add_verdict_block(38, "Digital Marketing Pro",
    "STUDY", "P3",
    "Impressive claims (115 commands, 25 agents, 64 scripts, 67 MCP integrations), but 14 stars on a repo claiming this scope is a red flag. Either quality does not match claims or it is untested. Worth a deeper technical review of actual script quality, but do not adopt blindly.",
    "If creative fatigue prediction scripts and self healing campaign logic are functional, adapt for creative analyst agent.",
    "Massive overlap with OFM's entire infrastructure. Competing architecture, not complement.")

add_verdict_block(39, "Growth Marketing Agents",
    "ADOPT", "P2",
    "Fills a genuine gap. OFM has no competitive intelligence agents. Three agents (Website Intelligence, Meta Ads Library, Google Ads Transparency) provide structured competitor analysis OFM currently does manually or not at all. Uses public APIs and data sources.",
    "Add as new agents or skills. Meta Ads Library supports creative audits (understand competitor creative before recommending direction). Google Ads Transparency supports keyword strategy. Website Intelligence feeds landing page audits.",
    "None. Purely additive. Meta Ads Library analysis is different from Meta Ads MCP (client accounts vs competitor research).")

add_verdict_block(40, "Claude MCP Marketing",
    "SKIP", "N/A",
    "Beginner oriented installer for basic MCP servers (Brave Search, Puppeteer, Slack). OFM's MCP stack is far more sophisticated. 10 stars.",
    "N/A.",
    "Puppeteer overlaps with Chrome MCP. Everything else irrelevant.")

# -- Category 6 --
add_heading_styled("3.6 Development & Browser Automation", 2)

add_verdict_block(41, "Context7",
    "ADOPT", "P1",
    "48.4K stars. Injects current, version specific documentation for any library directly into context. When building GTM tags, SuiteCommerce extensions (Backbone.js/RequireJS), Shopify modifications, or Klaviyo integrations, Claude's training data may reference outdated APIs. Context7 fixes this. Setup is trivial (npx ctx7 setup).",
    "Immediate value for SuiteCommerce development, Shopify Liquid, WooCommerce hooks, Klaviyo API, GA4 Measurement Protocol, and GTM Custom HTML JavaScript. Add 'use context7' to relevant skill files.",
    "None. Purely additive.")

add_verdict_block(42, "Repomix",
    "STUDY", "P3",
    "22.4K stars. Packs entire repo into single AI friendly file with token counting and tree sitter compression. Main use case for OFM would be onboarding new client codebases (SuiteCommerce themes, Shopify repos) for tracking integration analysis.",
    "Useful when onboarding new ecommerce clients. Not a daily tool.",
    "None. Standalone CLI.")

add_verdict_block(43, "UI/UX Pro Max Skill",
    "SKIP", "N/A",
    "Design system generator for building applications (React, Next.js, Vue, SwiftUI). OFM does not build applications. 39.6K stars driven by the dev community, zero relevance to paid media.",
    "N/A.",
    "Tangential to landing page analysis but fundamentally different scope.")

add_verdict_block(44, "Obsidian Skills",
    "SKIP", "N/A",
    "OFM does not use Obsidian. Knowledge management handled through file based client memory and Asana.",
    "N/A.",
    "Would compete with OFM's existing memory system.")

add_verdict_block(45, "n8n MCP",
    "SKIP", "N/A",
    "OFM uses Zapier, not n8n. Adopting would require migrating automation infrastructure. If OFM ever switches to n8n, this becomes P1 instantly.",
    "N/A unless platform migration.",
    "Competes with Zapier MCP integration.")

add_verdict_block(46, "Playwright MCP (Microsoft)",
    "STUDY", "P3",
    "28.6K stars, Microsoft backed. Offers headless execution, multi browser testing, and device emulation that Chrome MCP does not provide. Potential for automated tracking QA (headless browser loads pages, verifies tag firing across browsers) without needing a visible Chrome window.",
    "Automated tracking verification as OFM scales QA processes. Not urgent but worth evaluating.",
    "Partial overlap with Chrome MCP. Different enough to coexist (accessibility tree vs visual, headless vs interactive).")

add_verdict_block(47, "Playwright MCP (ExecuteAutomation)",
    "SKIP", "N/A",
    "If OFM adopted Playwright, Microsoft's official version is the clear choice. No reason to evaluate two Playwright MCPs.",
    "N/A.",
    "Redundant with both Chrome MCP and Microsoft Playwright.")

add_verdict_block(48, "GitHub MCP",
    "STUDY", "P2",
    "27.7K stars, official GitHub product. Would add native MCP access to repo management, issue/PR handling, Actions monitoring directly in Claude's tool loop without shelling out to gh. Value depends on volume of GitHub work.",
    "Streamline internal development workflow: issues for skill improvements, PRs for infrastructure, CI/CD monitoring.",
    "Overlaps with gh CLI via Bash. Adds convenience, not new capabilities.")

add_verdict_block(49, "Claude Code MCP (steipete)",
    "SKIP", "N/A",
    "Runs Claude Code with --dangerously-skip-permissions, bypassing all safety checks. OFM's security rules explicitly prohibit this. The entire design philosophy (bypass safety, execute autonomously) conflicts with OFM's 'never touch client systems without approval' mandate.",
    "N/A. Security model fundamentally incompatible.",
    "Directly conflicts with .claude/rules/security.md and CLAUDE.md safety rules.")

doc.add_page_break()

# ============================================================
# BATCH 1: TOP 10 POWER STACK
# ============================================================
add_heading_styled("4. Batch 1: Top 10 Power Stack", 1)

doc.add_paragraph("Install in this order. Each builds on the previous.")

add_table(
    ["Rank", "Repo", "Action", "Why First"],
    [
        ["1", "Context7 (#41)", "Install", "Fastest win. npx ctx7 setup. Current docs for every library OFM touches."],
        ["2", "Anthropic Official Skills (#19)", "Install", "Native DOCX/PPTX/XLSX/PDF creation for all OFM deliverables."],
        ["3", "Hooks Mastery (#13)", "Study + Build", "Unlocks 11 untapped hook events. Build SessionStart, Stop, PermissionRequest hooks."],
        ["4", "Meta Ads Analyzer (#33)", "Extract refs", "9 reference docs on Meta mechanics. Extract into skills/meta-ads-tracking/references/."],
        ["5", "Everything Claude Code (#9)", "Extract patterns", "Build instinct learning, hook profiles, and checkpoint to Asana from ECC's architecture."],
        ["6", "Claude Ads (#26)", "Gap analysis", "Cross reference 190 audit checks against OFM's existing skills. Evaluate A-F scoring."],
        ["7", "Growth Marketing Agents (#39)", "Adapt + Install", "Competitive intelligence agents for Meta Ads Library and Google Ads Transparency."],
        ["8", "GSD (#8)", "Extract pattern", "Phase based workflow for complex multi step engagements. Fresh context per wave."],
        ["9", "Context Mode (#2)", "Extract pattern", "MCP response sandboxing for verbose tool outputs during multi account audits."],
        ["10", "Best Practice (#18)", "Reference", "Implement 50% compact discipline. Evaluate CLAUDE.md sizing."],
    ]
)

doc.add_page_break()

# ============================================================
# BATCH 1: ARCHITECTURE BLUEPRINT
# ============================================================
add_heading_styled("5. Batch 1: Architecture Blueprint", 1)

doc.add_paragraph("Six layer architecture showing how adopted repos integrate with OFM's existing infrastructure.")

layers = [
    ("Layer 6: Output", [
        "Anthropic Document Skills (DOCX, PPTX, XLSX, PDF) [NEW #19]",
        "Weighted A-F Audit Scoring [BUILD CUSTOM from #26]",
        "Client memory with progressive disclosure [ENHANCE from #3]",
    ]),
    ("Layer 5: Hooks & Safety", [
        "deny-platform-writes (PreToolUse) [EXISTING]",
        "pre-compact-preserve (PreCompact) [UPGRADE from #2]",
        "SessionStart auto-loader [BUILD NEW from #13]",
        "Stop session summary [BUILD NEW from #13]",
        "PermissionRequest auto-approve reads [BUILD NEW from #13]",
        "Skill auto-suggest (UserPromptSubmit) [BUILD NEW from #11]",
        "Hook profiles (minimal/standard/strict) [BUILD NEW from #9]",
    ]),
    ("Layer 4: Orchestration", [
        "CLAUDE.md orchestration rules [EXISTING]",
        "Swarm mode parallel execution [EXISTING]",
        "Phase-based complex work lifecycle [BUILD NEW from #8]",
        "Instinct learning pipeline [BUILD NEW from #9]",
        "Quality gates at stage transitions [BUILD NEW from #25]",
    ]),
    ("Layer 3: Agents", [
        "9 existing domain agents [EXISTING]",
        "Competitive Intelligence agents [NEW #39]",
        "Three-tier model routing (Opus/Sonnet/Haiku) [EVALUATE from #14]",
    ]),
    ("Layer 2: Knowledge", [
        "CLAUDE.md + rules/ behavioral framework [EXISTING]",
        "27 custom paid media skills [EXISTING]",
        "Meta Ads reference docs [NEW #33]",
        "Claude Ads 190-check inventory [REFERENCE #26]",
        "Context7 library documentation [NEW #41]",
        "Anthropic skill spec standard [REFERENCE #19]",
    ]),
    ("Layer 1: MCP Data", [
        "Google Ads (cohnen) [EXISTING]",
        "Meta Ads (Pipeboard) [EXISTING]",
        "GA4 [EXISTING]",
        "GSC [EXISTING]",
        "Chrome [EXISTING]",
        "Asana [EXISTING]",
        "Context7 [NEW #41]",
        "Zapier [EXISTING]",
    ]),
]

for layer_name, items in layers:
    p = doc.add_paragraph()
    run = p.add_run(layer_name)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()

doc.add_page_break()

# ============================================================
# BATCH 1: CUSTOM BUILD RECOMMENDATIONS
# ============================================================
add_heading_styled("6. Batch 1: Custom Build Recommendations", 1)

b1_builds = [
    ("1. MCP Response Sandboxing (P1)", "Wrapper around verbose MCP tool calls (Google Ads GAQL, Meta insights with daily breakdowns, GA4 reports across 34 properties). Process full payloads externally, inject only actionable summaries into context. Critical for multi account audit sessions where raw data eats 50K+ tokens."),
    ("2. Instinct Based Learning Pipeline (P1)", "Formalize CLAUDE.md's self-learning protocol. When a session discovers something new (Meta changed CAPI dedup logic, a GTM pattern solves a consent mode edge case), capture as scored instinct in .claude/memory/instincts.md with confidence level, source, date. Periodically evolve high confidence instincts into skill updates."),
    ("3. Hook Profile System (P1)", "Environment variable (OFM_HOOK_PROFILE=minimal|standard|strict) controls hook intensity. Minimal: only deny-platform-writes. Standard: current 6 hooks. Strict: all hooks plus validation for tracking implementations. Reduces overhead on simple reads, maximizes protection on complex work."),
    ("4. Weighted Audit Scoring Engine (P2)", "A through F health score with severity weighted checks per platform, vertical specific weights (call tracking matters more for legal, feed quality matters more for ecommerce), and trend tracking across audit cycles."),
    ("5. Strategy Persistence Layer (P2)", "clients/[name]/strategy.md files that preserve strategic decisions (bid strategy rationale, audience architecture decisions, budget allocation logic) across sessions. Agents reference these before making recommendations."),
    ("6. Enhanced Pre-Compact Snapshot (P2)", "Upgrade pre-compact-preserve.js to generate structured snapshot with 10+ categories: active client, current task phase, findings so far, decisions made, files modified, Asana GIDs in play, MCP servers used, open questions, next steps."),
    ("7. Phase Based Complex Work Skill (P2)", "Discuss/Plan/Execute/Verify lifecycle with fresh context per execution wave. Scope: full tracking implementations, cross platform audits, new client onboarding."),
    ("8. Competitive Intelligence Skill (P3)", "Extend Growth Marketing Agents (#39) with OFM-specific analysis frameworks tied to client verticals."),
    ("9. LinkedIn & Microsoft Ads Skills (P3)", "If either platform enters OFM's client scope, build from Claude Ads (#26) check inventory rather than from scratch."),
    ("10. Automated Tracking QA Pipeline (P3)", "Headless browser loads client pages, verifies tag firing, event parameters, and conversion tracking across browsers."),
]

for title, desc in b1_builds:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    doc.add_paragraph(desc)

doc.add_page_break()

# ============================================================
# BATCH 2: INDIVIDUAL REVIEWS
# ============================================================
add_heading_styled("7. Batch 2: Individual Reviews (33 items)", 1)

# -- B2 Category 1 --
add_heading_styled("7.1 Project Execution & Workflow Engines", 2)

add_verdict_block(1, "GSD (gsd-build)",
    "STUDY", "P2",
    "This IS the evolved version of the glittercowboy/get-shit-done from Batch 1 (org transfer, 301 redirect confirmed, 27.9K stars). The fresh context per execution wave and atomic git commits pattern is strong. However, OFM's orchestration model already works: Michael prompts, Claude orchestrates swarm agents. GSD's interview/spec/execute flow adds ceremony OFM does not need for most client tasks.",
    "Study the 'fresh context per plan' pattern for large multi-phase projects like CWV remediation on Verocious Motorsports. The /gsd:verify-work pattern could strengthen OFM's audit verification workflow.",
    "Supersedes Batch 1 #8 (GSD glittercowboy). Overlaps conceptually with OFM's existing orchestration model.")

add_verdict_block(2, "Planning with Files (Manus Style)",
    "STUDY", "P2",
    "The filesystem as persistent working memory concept is solid. The three file pattern (task_plan.md, findings.md, progress.md) and the '2 Action Rule' (save findings after every 2 view/browser operations) address real context loss problems. The PreToolUse hook to re-read plans before decisions is clever. However, adding hooks that fire on every tool use adds overhead to every single operation.",
    "The 2 Action Rule pattern should be absorbed into OFM's CLAUDE.md for Chrome browser sessions where findings get lost during long navigation sequences. The three file pattern is useful for multi-session projects like client audits spanning multiple days.",
    "Overlaps with Batch 1 #1 (Continuous Claude v3) and #2 (Context Mode) on session persistence. Complementary rather than conflicting.")

# -- B2 Category 2 --
add_heading_styled("7.2 Context Engineering & Architecture Knowledge", 2)

add_verdict_block(3, "Agent Skills for Context Engineering",
    "STUDY", "P3",
    "Educational resource covering context fundamentals, degradation patterns, compression strategies, and multi-agent architectures. Useful for understanding WHY context breaks and how to prevent it. Not installable infrastructure, just knowledge. 7.8K stars suggests good content quality.",
    "One-time read to inform how OFM writes CLAUDE.md, structures agent prompts, and handles long sessions. Could improve how OFM structures its PreCompact hooks and context preservation rules.",
    "None. Reference material only.")

add_verdict_block(4, "Claude Code System Prompts (Piebald AI)",
    "ADOPT", "P1",
    "Complete documentation of Claude Code's internal system prompts across 122 versions. Understanding the exact system prompt, all 18 built-in tool descriptions, sub-agent prompts, and system reminders is invaluable for writing CLAUDE.md files and hooks that work WITH the system rather than against it. Updated within minutes of every Claude Code release.",
    "Use to optimize OFM's CLAUDE.md by understanding exactly what instructions Claude Code already has (avoid duplicating or contradicting built-in behavior). Use to write better hooks by understanding the exact lifecycle events and their payloads. Reference when debugging unexpected agent behavior.",
    "None. Pure reference material that enhances everything else.")

add_verdict_block(5, "Claude Code Best Practice (shanraisshan)",
    "STUDY", "P2",
    "Compiled best practices from power users. Key insights worth absorbing: manual /compact at max 50%, feature-specific subagents over general ones, ultrathink keyword for high-effort reasoning. The 'dumb zone' concept (agents lose effectiveness as context fills) validates OFM's existing swarm approach.",
    "Validate OFM's existing practices and identify any blind spots. The 50% compact threshold is a good rule to encode into OFM's workflow.",
    "Overlaps with Batch 1 #18 (same repo was listed in both batches). Skip if already reviewed in Batch 1.")

add_verdict_block(6, "Claude Code Guide (zebbern)",
    "SKIP", "N/A",
    "Beginner to power user guide. OFM is already well past this level of sophistication with 27 skills, 9 agents, 6 hooks, and 6 MCP servers. Nothing here that is not already documented in Anthropic's official docs or learned through OFM's direct experience.",
    "N/A.",
    "None.")

# -- B2 Category 3 --
add_heading_styled("7.3 Skill Generation & Management", 2)

add_verdict_block(7, "Skill Seekers",
    "STUDY", "P2",
    "Converts documentation websites, GitHub repos, and PDFs into structured Claude skills. 10.7K stars, active development. The concept is relevant for bootstrapping new platform skills from official docs. However, OFM's 27 existing skills are heavily customized with agency-specific audit workflows and scoring rubrics that no auto-generator would produce.",
    "Best used when OFM onboards a new platform or tool (e.g., a new ecommerce platform, a new ad channel). Point it at the docs, get a first-draft skill, then manually refine with OFM-specific audit logic and standards.",
    "None. Complementary to OFM's self-learning protocol.")

add_verdict_block(8, "Claude Skills Supercharged",
    "SKIP", "N/A",
    "Uses Haiku API calls on every prompt to auto-score and inject relevant skills. Dead project (37 stars, last updated January 2025). Adds API cost and latency overhead. OFM's keyword-based skill activation is sufficient for domain-specific skills where triggers are unambiguous.",
    "N/A.",
    "Would require restructuring OFM's hook system and adding API key management.")

add_verdict_block(9, "SkillForge",
    "SKIP", "N/A",
    "Impressive meta-skill generator with 4-phase pipeline and peer review, but designed for teams starting from zero. OFM already has 27 production skills and a self-learning protocol. The generated skills would need complete rewriting for OFM's domain-specific audit workflows.",
    "N/A.",
    "Redundant with OFM's existing self-learning protocol.")

add_verdict_block(10, "meta_skilld",
    "SKIP", "N/A",
    "Enterprise-grade Rust CLI for skill management with SQLite persistence, hybrid search, Thompson sampling, and MCP server mode. Massive overkill for a single-project, single-operator setup. Requires Rust 1.85+ compilation.",
    "N/A.",
    "None, but unnecessary complexity.")

# -- B2 Category 4 --
add_heading_styled("7.4 Web Research & Scraping", 2)

add_verdict_block(11, "Firecrawl CLI + Skill",
    "ADOPT", "P2",
    "The core Firecrawl product (91K stars) is the most battle-tested web scraping infrastructure in the AI ecosystem. The Claude plugin wraps the CLI as a skill with scrape, crawl, search, map, and agent commands. Fills a genuine gap: bulk scraping, recursive site crawling, and clean markdown output that Chrome browser automation handles poorly.",
    "Bulk competitor landing page audits across dozens of URLs. Crawling client sites to map all URLs for tracking audit coverage. Scraping platform documentation to feed into skill generation. Web search + scrape for competitive research. Free tier (500 credits/month) covers light usage.",
    "Complements (not replaces) claude-in-chrome. Chrome for interactive platform UIs; Firecrawl for bulk scraping and research.")

add_verdict_block(12, "Dev Browser",
    "SKIP", "N/A",
    "Well-built browser automation skill (3.8K stars) but functionally equivalent to OFM's existing claude-in-chrome MCP. Chrome Extension mode controls existing logged-in browser, which is exactly what claude-in-chrome already does. 3 months stale.",
    "N/A.",
    "Direct overlap with claude-in-chrome MCP.")

add_verdict_block(13, "Google AI Mode Skill",
    "SKIP", "N/A",
    "Stealth headless browser that scrapes Google AI Mode responses. Fragile (relies on Google not detecting the bot), 2 months stale, and redundant. OFM already has WebSearch as a built-in tool and Chrome for manual Google navigation.",
    "N/A.",
    "Redundant with WebSearch tool and Chrome browser.")

add_verdict_block(14, "Playwright CLI + Skill",
    "STUDY", "P2",
    "Microsoft themselves recommend CLI + Skills over MCP for token efficiency. Playwright Skill (1.9K stars) lets Claude write and execute arbitrary automation scripts with screenshots and console output. Adds capabilities chrome-in-chrome does not have: automated visual regression testing, structured form submission testing, multi-viewport responsive testing.",
    "Self-QA on Webflow/Shopify builds, landing page audit automation, visual regression testing on client sites after changes, form testing for lead gen clients. Runs in separate Chromium instance, no interference with chrome-in-chrome.",
    "Runs alongside claude-in-chrome (complementary, not competing). Overlaps with Batch 1 #46 and #47.")

# -- B2 Category 5 --
add_heading_styled("7.5 Prompt Preprocessing & Input Quality", 2)

add_verdict_block(15, "Prompt Improver",
    "STUDY", "P2",
    "UserPromptSubmit hook that evaluates prompt clarity before execution. Clear prompts pass through with only 189 tokens overhead. Vague prompts trigger clarifying questions. Bypass prefixes (* to skip, / for commands, # for memorize). The pattern of pre-evaluating prompts is sound, though OFM's prompts tend to be precise since Michael is highly technical.",
    "The bypass prefix pattern is worth absorbing. The skill-based architecture with progressive disclosure references is a clean implementation to study. Could help when less experienced team members use the system.",
    "Overlaps with #16 and #17. Choose one approach, not all three.")

add_verdict_block(16, "Prompt Enhancer",
    "SKIP", "N/A",
    "More aggressive prompt transformation than the Improver. Auto-injects 'requirements, edge cases, performance considerations' without asking. This would be actively harmful for OFM's workflow where Michael issues precise, targeted instructions and does not want them silently rewritten into bloated engineering specs.",
    "N/A.",
    "Overlaps with #15 and #17. Would conflict with OFM's direct, concise workflow.")

add_verdict_block(17, "UserPromptSubmit Hook (veteranbv)",
    "ADOPT (adapt)", "P2",
    "Self-documenting prompt enhancement with 18 development mode flags. The key insight is not the specific flags but the extensible pattern: OFM could add agency-specific flags like -bp for Blessed Performance context, -meta for Meta Ads mode, -gads for Google Ads mode, -vms for Verocious Motorsports context. Auto-injects current date and git branch.",
    "Fork the pattern and build OFM-specific flags. Client flags auto-load client memory and set platform context. Mode flags configure audit depth and output format. Logging all interactions creates an audit trail.",
    "Overlaps with #15 and #16. This is the best of the three for OFM because it is extensible rather than opinionated.")

# -- B2 Category 6 --
add_heading_styled("7.6 Content Quality & Writing", 2)

add_verdict_block(18, "Humanizer",
    "STUDY", "P3",
    "Claude Code skill that removes AI writing patterns. 2.9K stars. Relevant because OFM already has strict writing rules (no hyphens/dashes, no em dashes, natural conversational business prose). However, absorbing the pattern rules into CLAUDE.md is better than running a separate skill post-hoc.",
    "Review the specific transformation rules and add any missing patterns to OFM's CLAUDE.md writing instructions. Better to prevent AI-sounding text at generation time than fix it after.",
    "Overlaps with #19 and #20. Complementary to existing CLAUDE.md writing rules.")

add_verdict_block(19, "HumanizerAI",
    "HARD SKIP", "N/A",
    "Sends client text to an external API for AI detection scoring and humanization. This is a data security risk for an agency handling client ad copy, campaign strategies, and business intelligence. The text goes through a third-party service with unknown data retention policies.",
    "N/A.",
    "Conflicts with OFM's security rules about never exposing client data to unauthorized external services.")

add_verdict_block(20, "Beautiful Prose",
    "STUDY", "P3",
    "Hard-edged writing style contract for forceful English prose without AI tics. The style enforcement approach is more aligned with OFM's needs than post-hoc humanization. The specific rules about what constitutes AI-sounding text are the valuable component.",
    "Extract specific anti-AI-tic rules and incorporate into OFM's CLAUDE.md writing instructions. One-time knowledge transfer, not an ongoing dependency.",
    "Overlaps with #18 and OFM's existing writing rules in CLAUDE.md.")

# -- B2 Category 7 --
add_heading_styled("7.7 Security & Configuration", 2)

add_verdict_block(21, "AgentShield",
    "ADOPT", "P1",
    "Security scanner with 1,282 tests, 98% coverage, 102 static analysis rules. Scans CLAUDE.md, settings.json, MCP configs, hooks, agent definitions, and skills. The --opus flag runs red team/blue team/auditor pipeline. Given that OFM handles client ad accounts, API tokens, and business data, a security audit of the Claude Code configuration is essential and overdue.",
    "Run immediately against OFM's current setup. Identify any secrets exposure in CLAUDE.md or skill files. Audit MCP server risk profiles. Verify hook injection resistance. Run periodically after configuration changes.",
    "None. Pure security tooling that complements everything.")

add_verdict_block(22, "Parry (Prompt Injection Scanner)",
    "MONITOR", "P3",
    "Prompt injection scanner for Claude Code hooks. Early development. Worth monitoring given OFM handles client ad accounts and API credentials, but not mature enough to adopt yet.",
    "Revisit when more mature. The concept could be built into OFM's PreToolUse hooks as a lightweight injection check.",
    "None.")

add_verdict_block(23, "Varlock",
    "STUDY", "P2",
    "Secure environment variable management ensuring secrets are never exposed in Claude sessions, terminals, logs, or git commits. Critical concept for agency work where you handle Google Ads API tokens, Meta access tokens, Stape credentials, and client-specific API keys.",
    "Review the enforcement mechanisms. If it adds meaningful protection beyond OFM's existing security rules (which rely on Claude following instructions), adopt the enforcement patterns.",
    "Complementary to OFM's existing security rules in .claude/rules/security.md.")

add_verdict_block(24, "Dippy",
    "MONITOR", "P2",
    "AST-based bash command parsing that auto-approves safe commands while prompting for destructive ones. Solves permission fatigue without disabling safety. However, with Auto Mode (#31) launching soon, this may be superseded by Anthropic's official solution.",
    "Evaluate against Auto Mode when it launches. If Auto Mode handles permission management well, Dippy becomes unnecessary.",
    "Potentially superseded by Auto Mode (#31).")

# -- B2 Category 8 --
add_heading_styled("7.8 Monitoring, Directories & Ecosystem Tracking", 2)

add_verdict_block(25, "Awesome Claude Plugins (quemsah)",
    "SKIP", "N/A",
    "Automated adoption metrics across 7,413 repos. Useful as a discovery tool but not something to install. OFM has already done its discovery work across both batches.",
    "N/A.",
    "None.")

add_verdict_block(26, "Awesome Claude Code (jqueryscript)",
    "SKIP", "N/A",
    "Another curated list. OFM has already identified and evaluated the relevant repos through these two batch audits. Bookmark for monthly ecosystem scanning, no install.",
    "N/A.",
    "Overlaps with Batch 1 #10.")

add_verdict_block(27, "Claude Agent Blueprints",
    "SKIP", "N/A",
    "Index of 75+ Claude Code repositories. The 'Agent Workspace Model' concept (Git repo as complete workspace for any activity) is already what OFM is doing. No new insights beyond what OFM has already implemented.",
    "N/A.",
    "None.")

add_verdict_block(28, "Build with Claude",
    "SKIP", "N/A",
    "Marketplace directory. Useful for one-time discovery, not an ongoing tool. OFM has already performed thorough discovery across 82 items in two batches.",
    "N/A.",
    "None.")

# -- B2 Category 9 --
add_heading_styled("7.9 Social Media & Publishing", 2)

add_verdict_block(29, "Typefully",
    "SKIP", "N/A",
    "Social media scheduling across X, LinkedIn, Threads, Bluesky, Mastodon. OFM is a paid media agency, not a social media management agency. Social media publishing is not in the core service stack and would be scope creep.",
    "N/A.",
    "None.")

add_verdict_block(30, "X Article Publisher",
    "SKIP", "N/A",
    "Publishes markdown to X Articles. Same reasoning as #29. Not in OFM's service stack.",
    "N/A.",
    "None.")

# -- B2 Category 10 --
add_heading_styled("7.10 Upcoming / Just Announced", 2)

add_verdict_block(31, "Claude Code Auto Mode",
    "ADOPT", "P1 (when available)",
    "Removes the single biggest friction point in Claude Code workflows: constant permission prompts during development sessions. For OFM's swarm agent operations, permission fatigue across multiple parallel agents is a real productivity drag. Auto Mode would let read-only operations proceed without interruption.",
    "CRITICAL: Must verify that OFM's PreToolUse hooks (particularly deny-platform-writes.js) still fire under Auto Mode. If Auto Mode bypasses hooks, it would break OFM's safety guardrails. Test in isolated session before deploying to production workflow.",
    "May supersede Dippy (#24). Must coexist with OFM's security hooks.")

add_verdict_block(32, "Code Review",
    "SKIP", "N/A",
    "Multi-agent PR review tool. Designed for Team and Enterprise plans. OFM is a single-operator setup where Michael reviews his own code. The multi-agent PR review adds value in team environments, not solo developer workflows.",
    "N/A until OFM scales to multiple developers.",
    "None.")

add_verdict_block(33, "/loop Command",
    "ADOPT", "P1 (when available)",
    "Cron-style scheduling for continuous monitoring. Up to 50 concurrent tasks per session. This directly enables automated implementation of OFM's anomaly-flagging.md rules: hourly Google Ads budget pacing checks, 4-hour Meta CPM spike detection, daily cross-platform attribution reconciliation.",
    "Immediate use cases: /loop 1h scan Google Ads for limited budget campaigns hitting targets; /loop 4h check Meta CAPI event match quality across all accounts; /loop 1d run cross-platform conversion reconciliation. Aligns perfectly with OFM's existing anomaly-flagging rules.",
    "None. Enhances every monitoring and alerting workflow.")

doc.add_page_break()

# ============================================================
# BATCH 2: TOP 5
# ============================================================
add_heading_styled("8. Batch 2: Top 5", 1)

add_table(
    ["Rank", "Item", "Why Install First"],
    [
        ["1", "Auto Mode (#31)", "Removes permission friction across all workflows; must verify hook compatibility first"],
        ["2", "/loop Command (#33)", "Enables automated anomaly monitoring aligned to OFM's existing flagging rules"],
        ["3", "AgentShield (#21)", "Security audit of OFM's configuration is overdue; one-time scan identifies exposure risks"],
        ["4", "Claude Code System Prompts (#4)", "Understanding Claude Code's internals improves every CLAUDE.md instruction and hook"],
        ["5", "Firecrawl (#11)", "Fills the bulk scraping gap for competitor research and client site auditing"],
    ]
)

doc.add_page_break()

# ============================================================
# COMBINED POWER STACK
# ============================================================
add_heading_styled("9. Combined Power Stack (15 repos, both batches)", 1)

doc.add_paragraph("The ideal 15-item stack for OFM, in install order, pulling from both batches.")

add_table(
    ["Order", "Repo", "Batch", "Layer", "Why"],
    [
        ["1", "Auto Mode", "B2 #31", "Orchestration", "Removes permission friction; test hook compatibility first"],
        ["2", "/loop Command", "B2 #33", "Monitoring", "Enables continuous anomaly monitoring via existing flagging rules"],
        ["3", "AgentShield", "B2 #21", "Security", "One-time security audit, then periodic re-scans after config changes"],
        ["4", "System Prompts", "B2 #4", "Knowledge", "Reference for optimizing CLAUDE.md and hooks against actual system behavior"],
        ["5", "Context7", "B1 #41", "Knowledge", "Version-specific docs for Google Ads API, Meta Marketing API, etc."],
        ["6", "Official Skills", "B1 #19", "Skills", "Official document creation skills (DOCX, PPTX, PDF, XLSX) for deliverables"],
        ["7", "Hooks Mastery", "B1 #13", "Hooks", "All 13 hook events with working code; reference for building OFM hooks"],
        ["8", "Meta Ads Analyzer", "B1 #33", "Skills", "Breakdown Effect framework and Learning Phase diagnostics for Meta audits"],
        ["9", "UserPromptSubmit", "B2 #17", "Hooks", "Extensible flag system for client context and audit mode switching"],
        ["10", "Firecrawl", "B2 #11", "Research", "Bulk scraping for competitor analysis, client site crawling, doc extraction"],
        ["11", "Growth Mktg Agents", "B1 #39", "Agents", "Competitive intelligence for Meta Ad Library and Google Transparency"],
        ["12", "Planning w/ Files", "B2 #2", "Context", "Three-file working memory pattern for multi-session projects"],
        ["13", "Playwright Skill", "B2 #14", "QA/Testing", "Token-efficient browser testing for landing page audits and form validation"],
        ["14", "Varlock", "B2 #23", "Security", "Secret enforcement beyond CLAUDE.md honor-system rules"],
        ["15", "Skill Seekers", "B2 #7", "Skills", "Bootstrap skills from docs when onboarding new platforms"],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Already Adopted (do not re-install):")
run.bold = True
doc.add_paragraph("Google Ads MCP (cohnen) from B1 #28", style='List Bullet')
doc.add_paragraph("Meta Ads MCP (Pipeboard) from B1 #32", style='List Bullet')
doc.add_paragraph("GA4 MCP from B1 #35", style='List Bullet')
doc.add_paragraph("GSC MCP from B1 #36", style='List Bullet')

doc.add_page_break()

# ============================================================
# REDUNDANCY REPORT
# ============================================================
add_heading_styled("10. Redundancy Report", 1)

doc.add_paragraph("Which Batch 1 repos are superseded or changed by Batch 2 findings.")

add_table(
    ["Batch 1 Repo", "Status", "Reason"],
    [
        ["GSD (glittercowboy) B1 #8", "SUPERSEDED", "GSD (gsd-build) B2 #1 is the same repo after org transfer. 301 redirect confirmed."],
        ["Best Practice B1 #18", "DUPLICATE", "Same repo listed in both batches (B1 #18 and B2 #5). Already evaluated."],
        ["Claude Code Guide B1 #17", "DUPLICATE", "Same repo listed in both batches (B1 #17 and B2 #6). Already skipped."],
        ["Playwright MCP (MS) B1 #46", "PARTIALLY SUPERSEDED", "Microsoft recommends CLI+Skill (B2 #14) over MCP for token efficiency."],
        ["Playwright MCP (EA) B1 #47", "SUPERSEDED", "Between Microsoft official MCP and CLI+Skill, this adds no unique value."],
        ["Continuous Claude v3 B1 #1", "NOT superseded", "Planning with Files (B2 #2) is complementary (working memory vs session handoff)."],
        ["Context Mode B1 #2", "NOT superseded", "Focuses on post-compaction recovery; Planning with Files on within-session memory."],
        ["SuperClaude B1 #6", "Remains SKIP", "Nothing in Batch 2 changes this assessment."],
        ["Everything Claude Code B1 #9", "AgentShield extracted", "AgentShield (B2 #21) is standalone from this repo. Rest remains STUDY."],
    ]
)

doc.add_page_break()

# ============================================================
# BATCH 2: CUSTOM BUILD RECOMMENDATIONS
# ============================================================
add_heading_styled("11. Custom Build Recommendations (Batch 2)", 1)

doc.add_paragraph("Gaps where no existing repo serves OFM's exact needs. Build these in-house.")

b2_builds = [
    ("1. Client Context Flag System", "P1",
     "Build OFM-specific UserPromptSubmit hook with flags: -bp (Blessed Performance), -vms (Verocious Motorsports), -aus (Austenitex), -meta (Meta Ads audit mode), -gads (Google Ads audit mode), -tt (TikTok audit mode), -u (ultrathink 31,999 token budget), -audit (full audit mode). Fork veteranbv's pattern, strip generic flags, replace with OFM-specific flags."),
    ("2. Stape SGTM Health Monitor", "P1",
     "Build a /loop-compatible monitoring skill: check Stape container status via API, monitor CAPI event delivery rates per client, alert when event_id dedup rates drop below threshold, compare client-side vs server-side event counts, flag click ID capture failures. No existing repo covers Stape.io monitoring."),
    ("3. MCP Response Sandboxing Hook", "P2",
     "Build a PostToolUse hook: scan MCP responses for prompt injection attempts, validate ad account data stays associated with correct client, prevent cross-client data leakage in multi-account operations, log suspicious patterns. Parry (B2 #22) is too immature. Build custom with agency-specific threat models."),
    ("4. Weighted Audit Scoring Engine", "P2",
     "Build custom scoring: weight findings by OFM priority order (tracking > structure > bidding > audiences > creative > LP > attribution > wasted spend), calculate platform health scores using OFM-specific benchmarks, generate standard OFM deliverable format (Executive Summary, Score Snapshot, Issue Inventory, Detailed Findings, Action Plan)."),
    ("5. Cross-Platform Anomaly Dashboard", "P2",
     "Build /loop-compatible monitoring scripts aligned to anomaly-flagging.md: Hourly (Google Ads budget pacing, bidding learning status), Every 4 hours (Meta CPM spikes, creative frequency, CAPI match quality), Daily (cross-platform conversion reconciliation, GA4 event volumes), Weekly (Quality Score tracking, audience overlap, creative freshness). Outputs to Asana tasks on anomaly detection."),
    ("6. Klaviyo/ActiveCampaign Integration Skill", "P3",
     "Audit email/SMS flow alignment with paid media campaigns, check suppression list sync between ad platforms and email platforms, verify post-purchase flow triggers match conversion events, map customer journey touchpoints across paid + owned channels."),
    ("7. NetSuite/SuiteCommerce Tracking Skill", "P3",
     "Backbone.js AMD module architecture for tracking extensions, SC order model data layer extraction patterns, Pacejet integration points, multi-website tracking isolation (VMS Website ID 3, AUS Website ID 4 on shared account 606473)."),
]

for title, priority, desc in b2_builds:
    p = doc.add_paragraph()
    run = p.add_run(f"{title} ({priority})")
    run.bold = True
    run.font.size = Pt(10)
    doc.add_paragraph(desc)

doc.add_page_break()

# ============================================================
# ACTION PLAN
# ============================================================
add_heading_styled("12. Implementation Action Plan", 1)

phases = [
    ("Week 1 (Immediate)", [
        "Run AgentShield scan against current OFM setup. Fix any findings.",
        "Read Claude Code System Prompts repo. Cross-reference against current CLAUDE.md to identify instructions that duplicate or contradict built-in behavior. Optimize.",
        "When Auto Mode becomes available: test in isolated session. Verify PreToolUse hooks fire. If safe, enable for production workflows.",
        "When /loop becomes available: implement 3 monitoring loops aligned to anomaly-flagging.md.",
    ]),
    ("Week 2 to 3 (Near Term)", [
        "Build Client Context Flag System (Custom Build #1) using veteranbv's UserPromptSubmit pattern.",
        "Install Firecrawl plugin. Test with competitor landing page audit workflow.",
        "Absorb writing patterns from Humanizer (#18) and Beautiful Prose (#20) into CLAUDE.md. Do not install as separate skills.",
        "Review Varlock enforcement mechanisms. Adopt any that add protection beyond current security rules.",
    ]),
    ("Month 2 (Medium Term)", [
        "Build Stape SGTM Health Monitor (Custom Build #2) once /loop is stable.",
        "Build Weighted Audit Scoring Engine (Custom Build #4) to standardize audit deliverables.",
        "Study Planning with Files pattern. Implement the 2 Action Rule for Chrome browser sessions.",
        "Evaluate Playwright Skill for landing page QA automation on first relevant client project.",
    ]),
    ("Month 3+ (Backlog)", [
        "Build MCP Response Sandboxing Hook (Custom Build #3).",
        "Build Cross-Platform Anomaly Dashboard (Custom Build #5) as /loop matures.",
        "Build Klaviyo/ActiveCampaign Integration Skill (Custom Build #6) when next relevant client audit surfaces.",
        "Build NetSuite/SuiteCommerce Tracking Skill (Custom Build #7) for Verocious/Austenitex.",
    ]),
]

for phase_name, items in phases:
    p = doc.add_paragraph()
    run = p.add_run(phase_name)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()

# ============================================================
# FINAL NOTE
# ============================================================
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Bottom Line: ")
run.bold = True
run.font.size = Pt(11)
p.add_run(
    "Out of 82 items evaluated, 10 are worth adopting, 2 are already in production, and 31 are worth studying for pattern extraction. "
    "The biggest competitive advantage comes not from installing more repos but from the 7 custom builds that address the specific intersection "
    "of paid media agency operations, Stape.io server-side tracking, multi-platform MCP orchestration, and the client memory protocol OFM has already built. "
    "No open source repo addresses this intersection. That is where the real differentiation lives."
)

# ============================================================
# SAVE
# ============================================================
output_path = r"C:\Users\mtate\OneDrive\Desktop\ClaudeCode\Paid Media & Tracking Expert\OFM-Claude-Code-Repo-Audit-Complete.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
